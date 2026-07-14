"""
DOCX resume parser. Returns the same ParsedDocument shape as pdf_parser so
downstream code (skill extraction, ATS checks) doesn't need to care which
file type it originally was.
"""
import docx

from app.parsers.pdf_parser import ParsedDocument


def parse_docx(file_path: str) -> ParsedDocument:
    try:
        document = docx.Document(file_path)

        paragraphs = [p.text for p in document.paragraphs]
        # Tables are a common ATS-parseability failure point in DOCX resumes —
        # extract their text too so skill extraction still works, but we track
        # table presence separately for the ATS checker.
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_text.append(cell.text)

        raw_text = "\n".join(paragraphs + table_text).strip()
        char_count = len(raw_text)

        has_images = False
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                has_images = True
                break

        non_ascii_count = sum(1 for c in raw_text if ord(c) > 127)
        non_ascii_ratio = (non_ascii_count / char_count) if char_count else 0.0

        return ParsedDocument(
            raw_text=raw_text,
            page_count=1,  # DOCX has no reliable page count without rendering
            has_extractable_text=char_count > 20,
            is_multi_column=len(document.tables) > 0,  # tables are the DOCX analogue of columns
            has_embedded_images=has_images,
            char_count=char_count,
            non_ascii_ratio=non_ascii_ratio,
            per_page_word_x_positions=[],
        )
    except Exception as exc:
        raise ValueError(f"Could not parse DOCX: {exc}") from exc
